import logging
from pathlib import Path
from PIL import Image
import pytesseract
import fitz  # PyMuPDF
from docx import Document
import openpyxl
from pdf2image import convert_from_path
from config import Config

logger = logging.getLogger(__name__)

class FileProcessor:
    def __init__(self):
        self.download_dir = Path(Config.DOWNLOAD_DIR)
        self.download_dir.mkdir(exist_ok=True)
    
    def process_image(self, file_path: Path) -> str:
        """OCR обработка изображений"""
        try:
            img = Image.open(file_path)
            if img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Оптимизация размера для OCR
            if max(img.size) > 3000:
                img.thumbnail((2500, 2500), Image.Resampling.LANCZOS)
            
            text = pytesseract.image_to_string(img, lang=Config.TESSERACT_LANG)
            return text.strip() if text else "📷 Текст на изображении не распознан"
            
        except Exception as e:
            logger.error(f"Image OCR error: {e}")
            return f"❌ Ошибка распознавания изображения: {e}"
    
    def process_pdf(self, file_path: Path) -> str:
        """Обработка PDF файлов"""
        try:
            doc = fitz.open(str(file_path))
            full_text = []
            ocr_pages = []
            
            # Извлечение текста из PDF
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                text = page.get_text().strip()
                if text:
                    full_text.append(f"📄 Страница {page_num + 1}:\n{text}")
                else:
                    ocr_pages.append(page_num)
            
            # OCR для страниц без текста
            if ocr_pages:
                full_text.append("🔍 Страницы требующие OCR: " + ", ".join(map(str, [p+1 for p in ocr_pages])))
                try:
                    images = convert_from_path(
                        str(file_path), 
                        first_page=min(ocr_pages)+1, 
                        last_page=max(ocr_pages)+1
                    )
                    for i, page_num in enumerate(ocr_pages):
                        if i < len(images):
                            txt = pytesseract.image_to_string(images[i], lang=Config.TESSERACT_LANG)
                            if txt.strip():
                                full_text.append(f"📄 Страница {page_num + 1} (OCR):\n{txt}")
                except Exception as ocr_error:
                    logger.error(f"PDF OCR failed: {ocr_error}")
            
            doc.close()
            return "\n\n".join(full_text) if full_text else "📄 В PDF не найден текст"
            
        except Exception as e:
            logger.error(f"PDF processing error: {e}")
            return f"❌ Ошибка обработки PDF: {e}"
    
    def process_docx(self, file_path: Path) -> str:
        """Обработка DOCX файлов"""
        try:
            doc = Document(str(file_path))
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Обработка таблиц
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        tables_text.append(row_text)
            
            result = []
            if paragraphs:
                result.append("📝 Текст документа:\n" + "\n".join(paragraphs))
            if tables_text:
                result.append("📊 Таблицы:\n" + "\n".join(tables_text))
            
            return "\n\n".join(result) if result else "📝 Документ пуст"
            
        except Exception as e:
            logger.error(f"DOCX processing error: {e}")
            return f"❌ Ошибка обработки DOCX: {e}"
    
    def process_xlsx(self, file_path: Path) -> str:
        """Обработка XLSX файлов"""
        try:
            wb = openpyxl.load_workbook(str(file_path), data_only=True)
            all_sheets_data = []
            
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                sheet_data = []
                
                # Читаем данные построчно
                for row in ws.iter_rows(values_only=True):
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    if any(cell.strip() for cell in row_data if cell):
                        sheet_data.append("\t".join(row_data))
                
                if sheet_data:
                    all_sheets_data.append(f"📊 Лист: {sheet_name}\n" + "\n".join(sheet_data))
            
            wb.close()
            return "\n\n" + "="*50 + "\n\n".join(all_sheets_data) if all_sheets_data else "📊 Файл не содержит данных"
            
        except Exception as e:
            logger.error(f"XLSX processing error: {e}")
            return f"❌ Ошибка обработки XLSX: {e}"
    
    def get_file_type(self, filename: str, mime_type: str) -> str:
        """Определение типа файла"""
        file_ext = Path(filename).suffix.lower() if filename else ""
        
        if mime_type.startswith('image/') or file_ext in ['.jpg', '.jpeg', '.png', '.bmp']:
            return 'image'
        elif mime_type == 'application/pdf' or file_ext == '.pdf':
            return 'pdf'
        elif 'wordprocessingml' in mime_type or file_ext in ['.docx', '.doc']:
            return 'docx'
        elif 'spreadsheetml' in mime_type or file_ext in ['.xlsx', '.xls']:
            return 'xlsx'
        else:
            return 'unknown'
